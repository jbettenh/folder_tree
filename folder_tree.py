#! python 3
# This will list out all folders and subfolders for a given directory.

import argparse
from pathlib import Path
from itertools import islice
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

# prefix components:
space =  '    '
branch = '│   '
# pointers:
tee =    '├── '
last =   '└── '


def arguments():
    parser = argparse.ArgumentParser(description='Parser')
    parser.add_argument('--path', type=Path,default=Path(__file__).absolute().parent / "Users", help='This is the path to the folder')

    return parser.parse_args()


def tree(dir_path: Path, level: int = -1, limit_to_directories: bool = False, length_limit: int = 1000):
        """Given a directory Path object print a visual tree structure"""
        dir_path = Path(dir_path)  # accept string coerceable to Path
        files = 0
        directories = 0

        def inner(dir_path: Path, prefix: str = '', level=-1):
            nonlocal files, directories
            if not level:
                return  # 0, stop iterating
            if limit_to_directories:
                contents = [d for d in dir_path.iterdir() if d.is_dir()]
            else:
                contents = list(dir_path.iterdir())
            pointers = [tee] * (len(contents) - 1) + [last]
            for pointer, path in zip(pointers, contents):
                if path.is_dir():
                    yield prefix + pointer + path.name
                    directories += 1
                    extension = branch if pointer == tee else space
                    yield from inner(path, prefix=prefix + extension, level=level - 1)
                elif not limit_to_directories:
                    yield prefix + pointer + path.name
                    files += 1

        print(dir_path.name)
        iterator = inner(dir_path, level=level)
        for line in islice(iterator, length_limit):
            print(line)
        if next(iterator, None):
            print(f'... length_limit, {length_limit}, reached, counted:')
        print(f'\n{directories} directories' + (f', {files} files' if files else ''))


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def main():
    args = arguments()

    for line in tree(args.path/ 'Documents', limit_to_directories=True):
        print(line)

   # for dirpath, dirnames, files in os.walk(args.path):
    #    path = dirpath.split('/')
     #   if dirnames:
      #      for folder in dirnames:
       #         print('-' + folder)

    document = docx.Document()
    p = document.add_paragraph('Folder: ')
    add_hyperlink(p, 'folder name', dirpath)
    document.save('test/demo_hyperlink.docx')


if __name__ == '__main__':
    main()